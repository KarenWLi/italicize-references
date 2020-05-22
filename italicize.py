import docx
import re

# Opening MS word document
doc = docx.Document('your-document')
problems = []

for p in doc.paragraphs:
    ## Because you can't just pull out a section, italicize it, and put it back in, 
    # we'll be breaking the citation into beginning, italics, and end
    # and using the p.add_run(text).italic = True command, add them back 
    beginning = p.text
    italics = ""
    end = "" 

    no_journal = True
    no_title = True
    #(There's such variation in titles that it might not be perfect; here are some issues to troubleshoot)

    #Case 1: It's a book
    #Then we have to italicize the title
    titles = re.finditer(r'\d{4}\.[^\n\d][a-zA-Z:,\'\-!\s]+\. ([A-Z]|\d)', p.text)
    
    for title in titles: #there should only ever be one, but just in case
        no_title = False
        s = title.start()
        e = title.end()

        beginning = p.text[0:s+5]
        italics = p.text[s+5:e-2]
        end = p.text[e-2 :len(p.text)-1]

        

            
    
    #Case 2: It's an article
    #Then we have to italicize the journal title 
    journals = re.finditer(r'" [a-zA-Z 0-9\-,:&]+ \d', p.text)

    editors = re.finditer(r'edited by', p.text)
    edited = re.findall(r'edited by', p.text)

    for journal in journals:
        no_journal = False
        s = journal.start()
        e = journal.end()

        if len(edited) == 0:   
            
            beginning = p.text[0:s+2]
            italics = p.text[s+2: e-2]
            end = p.text[e-2:len(p.text)-1]

        else:
            for editor in editors:
                t = editor.start()

                beginning = p.text[0:s+2]
                italics = p.text[s+2: t]
                end = p.text[t:len(p.text)-1]
    
    if no_journal and no_title:
        snippet = p.text[0:30] 
        problems.append(snippet)   

    p.text = ""
    p.add_run(beginning)
    p.add_run(italics).italic = True
    p.add_run(end)






print(problems)            
doc.save('Finished References.docx')